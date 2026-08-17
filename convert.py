#!/usr/bin/env python3
"""Matrix 工具箱 · 格式转换层

- excel_to_insert()  Excel + 建表语句 → INSERT INTO 语句（按 DDL 字段类型正确加引号）
- excel_to_pdf()     Excel → PDF（fpdf2 + macOS 中文字体，不依赖 Office/LibreOffice）
- excel_to_word()    Excel → Word（python-docx）
- image_to_rows()    图片 → 表格二维数组（macOS Vision OCR，按文字坐标还原行列）

选型说明：本机没有 LibreOffice/Excel/Numbers，所以 PDF/Word 都用纯 Python 生成；
OCR 走 macOS 原生 Vision（中文识别质量好、零第三方依赖），Swift 源码随包分发，
首次使用时就地编译并缓存。
"""

from __future__ import annotations

import json
import re
import subprocess
import sys
from datetime import date, datetime, time
from decimal import Decimal, InvalidOperation, localcontext
from pathlib import Path

import matrix_core as core

# ══════════════════════════════════════════════════════════════
# Excel → INSERT INTO
# ══════════════════════════════════════════════════════════════

# 需要加引号的类型；数值/布尔类型直接裸写
QUOTED_TYPES = ("char", "varchar", "string", "text", "date", "datetime",
                "timestamp", "time", "json", "binary", "blob")
NUMERIC_TYPES = ("int", "bigint", "smallint", "tinyint", "largeint", "decimal",
                 "double", "float", "numeric", "boolean", "bool")


def parse_ddl(ddl: str) -> dict:
    """从建表语句里解析出表名与字段（名称 + 类型）。

    兼容 StarRocks/MySQL 的 CREATE TABLE，容忍反引号、中文名、COMMENT。
    """
    text = core_strip_comments(ddl)
    m = re.search(r"CREATE\s+(?:EXTERNAL\s+)?TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?"
                  r"([`\w\u4e00-\u9fff.]+)\s*\(", text, re.I)
    if not m:
        raise ValueError("没找到 CREATE TABLE ... ( ，请贴完整的建表语句")
    table = m.group(1).replace("`", "")
    # 取最外层括号内的字段定义区
    start = m.end()
    depth, i = 1, start
    while i < len(text) and depth:
        if text[i] == "(":
            depth += 1
        elif text[i] == ")":
            depth -= 1
        i += 1
    body = text[start:i - 1]

    cols = []
    for line in _split_top_commas(body):
        line = line.strip()
        if not line:
            continue
        # 跳过表级约束/索引定义
        if re.match(r"(PRIMARY|UNIQUE|DUPLICATE|AGGREGATE)\s+KEY\s*\(", line, re.I):
            continue
        if re.match(r"(INDEX|KEY|CONSTRAINT|FOREIGN)\b", line, re.I):
            continue
        # 反引号/双引号中的字段名可包含空格、括号等字符。之前仅允许
        # ``\w + 中文``，会把 `金额（元）` 这类常见业务字段静默漏掉。
        cm = re.match(
            r'(?:`([^`]+)`|"([^"]+)"|([\w\u4e00-\u9fff]+))\s+'
            r'([A-Za-z][A-Za-z0-9_]*)',
            line,
        )
        if not cm:
            continue
        name = next(v for v in cm.groups()[:3] if v is not None)
        ctype = cm.group(4).lower()
        size_match = re.match(
            r"\s*\(\s*(\d+)(?:\s*,\s*(\d+))?\s*\)", line[cm.end():])
        cols.append({
            "name": name, "type": ctype,
            "quoted": _needs_quote(ctype),
            "nullable": not re.search(r"\bNOT\s+NULL\b", line, re.I),
            "precision": int(size_match.group(1)) if size_match else None,
            "scale": (int(size_match.group(2))
                      if size_match and size_match.group(2) is not None else None),
        })
    if not cols:
        raise ValueError("建表语句里没解析出任何字段")
    return {"table": table, "columns": cols}


def _needs_quote(ctype: str) -> bool:
    if any(ctype.startswith(t) for t in NUMERIC_TYPES):
        return False
    if any(t in ctype for t in QUOTED_TYPES):
        return True
    return True                      # 认不出的类型按字符串处理，更安全


def _split_top_commas(text: str) -> list[str]:
    """按最外层逗号切分，忽略括号内与字符串内的逗号"""
    out, depth, last, i = [], 0, 0, 0
    n = len(text)
    while i < n:
        c = text[i]
        if c in "'\"":
            q = c
            i += 1
            while i < n and not (text[i] == q and text[i - 1] != "\\"):
                i += 1
        elif c == "(":
            depth += 1
        elif c == ")":
            depth -= 1
        elif c == "," and depth == 0:
            out.append(text[last:i])
            last = i + 1
        i += 1
    out.append(text[last:])
    return out


def core_strip_comments(sql: str) -> str:
    """复用 SQL 层的注释清理，避免重复实现"""
    import sql_tools
    return sql_tools.strip_comments(sql)


def _sql_literal(value, col: dict) -> str:
    """把单元格值转成 SQL 字面量"""
    if value is None or (isinstance(value, str) and value.strip() == ""):
        return "NULL"
    if isinstance(value, bool):
        return "1" if value else "0"
    if isinstance(value, float) and abs(value) >= 1e15:
        raise ValueError(
            f"字段 `{col.get('name') or ''}` 是超过 15 位的 Excel 浮点数 {value!r}；"
            "Excel 可能已丢失末位精度，请把该列改为文本后重新导入"
        )
    ctype = str(col.get("type") or "").lower()
    if isinstance(value, datetime):
        if ctype == "date":
            s = value.strftime("%Y-%m-%d")
        elif ctype == "time":
            s = value.strftime("%H:%M:%S")
        else:
            s = value.strftime("%Y-%m-%d %H:%M:%S")
    elif isinstance(value, date):
        s = value.strftime("%Y-%m-%d")
    elif isinstance(value, time):
        s = value.strftime("%H:%M:%S")
    else:
        s = str(value).strip()
    # 空单元格已经在函数开头转成 SQL NULL。文本列里的真实字符串
    # "NULL" 必须原样保留，否则导入后无法区分文字与数据库空值。
    textual = any(ctype.startswith(t) for t in
                  ("char", "varchar", "string", "text", "json", "binary", "blob"))
    if s == "\\N" or (s.upper() == "NULL" and not textual):
        return "NULL"
    if not col["quoted"]:
        if ctype in ("boolean", "bool") and s.lower() in ("true", "false"):
            return "1" if s.lower() == "true" else "0"
        try:
            number_text = str(value).strip()
            if re.fullmatch(r"[+-]?\d{1,3}(?:,\d{3})+(?:\.\d+)?", number_text):
                number_text = number_text.replace(",", "")
            number = Decimal(number_text)
            if not number.is_finite():
                raise InvalidOperation
            scale = col.get("scale")
            if scale is not None:
                quantum = Decimal(1).scaleb(-int(scale))
                integer_digits = max(1, number.adjusted() + 1)
                with localcontext() as ctx:
                    ctx.prec = max(28, integer_digits + int(scale) + 4,
                                   len(number.as_tuple().digits) + int(scale) + 4)
                    quantized = number.quantize(quantum)
                precision = col.get("precision")
                if precision is not None:
                    integer_limit = int(precision) - int(scale)
                    quantized_integer_digits = (max(1, quantized.adjusted() + 1)
                                                if quantized else 1)
                    if quantized_integer_digits > integer_limit:
                        raise ValueError(
                            f"数值字段 `{col.get('name') or ''}` 的值 {s!r} "
                            f"超出 DECIMAL({precision},{scale}) 范围"
                        )
                return format(quantized, f".{int(scale)}f")
            # 无固定小数位的数值列：Excel 常把整数读成 100.0。
            if isinstance(value, float) and value.is_integer():
                return str(int(value))
            if isinstance(value, float):
                # openpyxl 可能返回 6172.799999999999 这类二进制浮点尾差。
                return format(value, ".15g")
            return format(number, "f")
        except ValueError:
            raise
        except (InvalidOperation, TypeError) as exc:
            raise ValueError(
                f"数值字段 `{col.get('name') or ''}` 收到非数值内容 {s!r}"
            ) from exc
    if isinstance(value, float) and value.is_integer():
        s = str(int(value))
    return "'" + s.replace("\\", "\\\\").replace("'", "\\'") + "'"


def excel_to_insert(xlsx: Path, ddl: str, batch: int = 500,
                    only_matched: bool = True, sheet=None) -> dict:
    """Excel + 建表语句 → INSERT 语句

    Excel 表头与 DDL 字段按名称匹配（忽略大小写与首尾空格）。
    """
    meta_ddl = parse_ddl(ddl)
    headers, rows, meta = core.read_sheet_meta(xlsx, sheet)
    if not headers:
        raise ValueError("没读到表头。" + "；".join(meta["warnings"])
                         + f"\n工作表清单: {[s['name'] for s in meta['sheets']]}")
    col_by_name = {c["name"].lower(): c for c in meta_ddl["columns"]}

    used, unmatched_header, positions = [], [], []
    matched_names: set[str] = set()
    duplicate_targets: list[str] = []
    for idx, h in enumerate(headers):
        key = str(h).strip().lower()
        c = col_by_name.get(key)
        if c:
            target_key = c["name"].lower()
            if target_key in matched_names:
                duplicate_targets.append(c["name"])
                continue
            matched_names.add(target_key)
            used.append(c)
            positions.append(idx)
        else:
            unmatched_header.append(h)
    if duplicate_targets:
        raise ValueError(
            "Excel 存在重复表头，会让 INSERT 重复写入同一目标字段："
            + "、".join(dict.fromkeys(duplicate_targets))
            + "。请先把表头改成唯一名称后重试")
    if not used:
        raise ValueError(f"Excel 表头与建表语句字段没有任何匹配。\n"
                         f"Excel 表头: {headers[:10]}\n"
                         f"DDL 字段: {[c['name'] for c in meta_ddl['columns']][:10]}")
    missing_cols = [c["name"] for c in meta_ddl["columns"]
                    if c["name"].lower() not in
                    {str(h).strip().lower() for h in headers}]

    col_list = ", ".join(f"`{c['name']}`" for c in used)
    statements, values_buf, skipped = [], [], 0
    for row_number, r in enumerate(rows, start=int(meta.get("header_row") or 1) + 1):
        if all(v is None or str(v).strip() == "" for v in r):
            skipped += 1
            continue
        vals = []
        for c, pos in zip(used, positions):
            v = r[pos] if pos < len(r) else None
            try:
                vals.append(_sql_literal(v, c))
            except ValueError as exc:
                raise ValueError(f"Excel 第 {row_number} 行：{exc}") from exc
        values_buf.append("(" + ", ".join(vals) + ")")
        if len(values_buf) >= batch:
            statements.append(_wrap_insert(meta_ddl["table"], col_list, values_buf))
            values_buf = []
    if values_buf:
        statements.append(_wrap_insert(meta_ddl["table"], col_list, values_buf))

    data_rows = len(rows) - skipped
    return {
        "table": meta_ddl["table"],
        "sql": "\n\n".join(statements),
        "stmt_count": len(statements),
        "row_count": data_rows,
        "skipped_empty": skipped,
        "matched_columns": [c["name"] for c in used],
        "unmatched_header": unmatched_header,
        "missing_columns": missing_cols,
        "batch": batch,
        "sheet": meta["sheet"],
        "sheets": [s["name"] for s in meta["sheets"]],
        "warnings": meta["warnings"],
    }


def _wrap_insert(table: str, col_list: str, values: list[str]) -> str:
    quoted_table = ".".join(
        "`" + part.replace("`", "``") + "`"
        for part in table.split(".") if part
    )
    if not quoted_table:
        raise ValueError("目标表名不能为空")
    return (f"INSERT INTO {quoted_table}\n  ({col_list})\nVALUES\n  "
            + ",\n  ".join(values) + ";")


# ══════════════════════════════════════════════════════════════
# Excel → PDF
# ══════════════════════════════════════════════════════════════

CJK_FONTS = (
    # Windows：fpdf2 优先微软雅黑；老版 pyfpdf 会自动跳过 TTC，改用 SimHei TTF。
    "C:/Windows/Fonts/msyh.ttc",
    "C:/Windows/Fonts/simhei.ttf",
    "C:/Windows/Fonts/msyh.ttf",
    # macOS
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    # 常见 Linux 桌面/容器
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf",
)

EMOJI_FONTS = (
    "C:/Windows/Fonts/seguiemj.ttf",
    "/System/Library/Fonts/Apple Color Emoji.ttc",
    "/usr/share/fonts/truetype/noto/NotoColorEmoji.ttf",
)


def _pick_font() -> str | None:
    for f in CJK_FONTS:
        if Path(f).exists():
            return f
    return None


def _prepare_font() -> tuple[str, bool]:
    """准备可用的中文字体，返回 (字体路径, 是否老版 pyfpdf 需要 uni=True)。

    老版 pyfpdf 用 uni=True 时会把解析缓存 .pkl 写到「字体文件所在目录」，
    而 /Library/Fonts 不可写，所以先把字体复制到自己的缓存目录再用。
    只接受 .ttf：老版 pyfpdf 不支持 .ttc 字体集合。
    """
    import inspect
    from fpdf import FPDF
    needs_uni = "uni" in inspect.signature(FPDF.add_font).parameters

    src = None
    for f in CJK_FONTS:
        if Path(f).exists() and (f.lower().endswith(".ttf") or not needs_uni):
            src = Path(f)
            break
    if src is None:
        raise RuntimeError(
            "找不到可用的中文字体。老版 pyfpdf 只支持 .ttf，"
            f"已尝试: {CJK_FONTS}")

    if not needs_uni:
        return str(src), False

    font_dir = core.CACHE_DIR / "fonts"
    font_dir.mkdir(parents=True, exist_ok=True)
    dst = font_dir / src.name
    if not dst.exists() or dst.stat().st_size != src.stat().st_size:
        import shutil
        # 用 copyfile 而非 copy2：系统字体带特殊 chflags，复制元数据会被拒
        shutil.copyfile(src, dst)
    return str(dst), True


def _display_value(value, header: str = "") -> str:
    """把工作簿值转成适合文档阅读的文本，而不是 Python 的原始表示。"""
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, datetime):
        if value.time() == time(0, 0):
            return value.strftime("%Y-%m-%d")
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, time):
        return value.strftime("%H:%M:%S")
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        h = str(header).lower()
        if any(k in h for k in ("百分比", "百分率", "占比", "比例")) or h.endswith("率"):
            return f"{value:.2%}"
        if any(k in h for k in ("金额", "价格", "单价", "费用", "余额", "收入", "支出", "总额", "税额")):
            return f"{value:,.2f}"
        if isinstance(value, int) or value.is_integer():
            return str(int(value))
        return format(value, ".15g")
    return str(value)


def _content_weight(value, header: str = "") -> int:
    text = _display_value(value, header)
    lines = text.replace("\r", "").split("\n")
    return max((len(line) + sum(1 for ch in line if ord(ch) > 127)
                for line in lines), default=4)


def _pdf_wrap(pdf, text: str, width: float, max_lines: int) -> tuple[list[str], bool]:
    """按实际字体宽度换行；返回 (行, 是否因安全上限被截断)。"""
    text = str(text).replace("\r", "").replace("\t", "    ")
    if not text:
        return [""], False
    limit = max(width - 2.0, 1.0)
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if paragraph == "":
            lines.append("")
            continue
        current = ""
        for ch in paragraph:
            candidate = current + ch
            if current and pdf.get_string_width(candidate) > limit:
                lines.append(current)
                current = ch
            else:
                current = candidate
        lines.append(current)
    clipped = len(lines) > max_lines
    if clipped:
        lines = lines[:max_lines]
        tail = lines[-1]
        while tail and pdf.get_string_width(tail + "…") > limit:
            tail = tail[:-1]
        lines[-1] = (tail + "…") if tail else "…"
    return lines or [""], clipped


def excel_to_pdf(xlsx: Path, out: Path, landscape: bool = True,
                 max_rows: int = 800, title: str = "", sheet=None) -> dict:
    """Excel → 可阅读的 PDF：跨平台中文字体、自动换行和重复表头。"""
    from fpdf import FPDF

    font_path, needs_uni = _prepare_font()
    headers, rows, meta = core.read_sheet_meta(xlsx, sheet)
    if not headers and not rows:
        raise ValueError("没读到任何内容。" + "；".join(meta["warnings"])
                         + f"\n工作表清单: {[s['name'] for s in meta['sheets']]}")
    total_rows = len(rows)
    truncated = total_rows > max_rows
    rows = rows[:max_rows]

    pdf = FPDF(orientation="L" if landscape else "P", unit="mm", format="A4")
    pdf.set_margins(8, 8, 8)
    pdf.set_auto_page_break(auto=False, margin=8)
    if needs_uni:
        pdf.add_font("cjk", "", font_path, uni=True)
    else:
        pdf.add_font("cjk", "", font_path)
    emoji_font = ""
    if hasattr(pdf, "set_fallback_fonts"):
        for candidate in EMOJI_FONTS:
            if not Path(candidate).exists():
                continue
            try:
                pdf.add_font("emoji", "", candidate)
                pdf.set_fallback_fonts(["emoji"])
                emoji_font = Path(candidate).name
                break
            except Exception:  # optional enhancement; CJK output must still work
                continue
    pdf.set_font("cjk", size=8)
    pdf.set_draw_color(190, 198, 210)
    pdf.set_line_width(0.2)

    def write_cell(width, height, text, align="L", next_line=False):
        if needs_uni:  # legacy PyFPDF compatibility
            pdf.cell(width, height, text, 0, 1 if next_line else 0, align)
        else:
            from fpdf.enums import XPos, YPos
            pdf.cell(width, height, text, border=0, align=align,
                     new_x=XPos.LMARGIN if next_line else XPos.RIGHT,
                     new_y=YPos.NEXT if next_line else YPos.TOP)

    ncol = max(len(headers), max((len(r) for r in rows), default=0)) or 1
    avail = (297 if landscape else 210) - 16
    weights = []
    for i in range(ncol):
        column_header = headers[i] if i < len(headers) else ""
        values = [column_header]
        values += [r[i] for r in rows[:200] if i < len(r)]
        weight = min(max(max((_content_weight(v, column_header) for v in values),
                             default=4), 4), 48)
        header_text = str(column_header).lower()
        if any(k in header_text for k in ("日期", "时间", "date", "time")):
            weight = max(weight, 12)
        if any(k in header_text for k in ("金额", "价格", "费用", "amount", "price")):
            weight = max(weight, 13)
        if any(k in header_text for k in ("百分", "比例")) or header_text.endswith("率"):
            weight = max(weight, 10)
        if any(k in header_text for k in ("电话", "手机", "联系方式")):
            weight = max(weight, 14)
        weights.append(weight)
    # A4 横向硬塞几十列会把每列压到几个毫米，最终逐字换行、每页只有
    # 三五行且大量截断。宽表改为“列组”分页，并在后续列组重复首列，
    # 让主键/编号始终可见，换取真正可读的输出。
    max_band_cols = 12 if landscape else 7
    if ncol <= max_band_cols:
        bands = [list(range(ncol))]
    else:
        payload_size = max_band_cols - 1
        bands = [[0, *range(start, min(start + payload_size, ncol))]
                 for start in range(1, ncol, payload_size)]

    def widths_for(indices):
        selected = [weights[i] for i in indices]
        total_weight = sum(selected) or 1
        widths = [max(avail * w / total_weight, 12) for w in selected]
        scale = avail / sum(widths)
        return [w * scale for w in widths]

    line_h = 4.2
    clipped_cells = 0
    clipped_headers = 0

    def layout(values, indices, widths, header_row=False):
        nonlocal clipped_cells, clipped_headers
        wrapped = []
        for original_i, width in zip(indices, widths):
            value = values[original_i] if original_i < len(values) else ""
            column_header = headers[original_i] if original_i < len(headers) else ""
            text = _display_value(value, column_header)
            lines, clipped = _pdf_wrap(pdf, text, width, 5 if header_row else 10)
            wrapped.append(lines)
            if clipped:
                if header_row:
                    clipped_headers += 1
                else:
                    clipped_cells += 1
        height = max(6.2, max(len(lines) for lines in wrapped) * line_h + 2.0)
        return wrapped, height

    def draw_layout(wrapped, height, indices, widths,
                    header_row=False, alternate=False):
        start_x, y = pdf.l_margin, pdf.get_y()
        if header_row:
            pdf.set_fill_color(47, 84, 150)
            pdf.set_text_color(255, 255, 255)
        else:
            color = 245 if alternate else 255
            pdf.set_fill_color(color, color + (1 if color < 255 else 0),
                               min(color + (3 if color < 255 else 0), 255))
            pdf.set_text_color(32, 39, 51)
        x = start_x
        for local_i, (original_i, width) in enumerate(zip(indices, widths)):
            pdf.rect(x, y, width, height, "DF")
            lines = wrapped[local_i]
            text_y = y + max((height - len(lines) * line_h) / 2, 0.8)
            align = "C" if header_row else "L"
            if not header_row and original_i < len(headers):
                h = str(headers[original_i])
                if any(k in h for k in ("日期", "时间", "状态", "比例", "百分")):
                    align = "C"
            for li, text in enumerate(lines):
                pdf.set_xy(x + 1, text_y + li * line_h)
                write_cell(max(width - 2, 0.1), line_h, text, align)
            x += width
        pdf.set_xy(start_x, y + height)
        pdf.set_text_color(32, 39, 51)

    def add_page(indices, widths, header_layout, band_no, first=False):
        pdf.add_page()
        if first and title:
            pdf.set_font("cjk", size=14)
            title_lines, _ = _pdf_wrap(pdf, title, avail, 2)
            for line in title_lines:
                write_cell(avail, 7, line, "C", next_line=True)
            pdf.ln(2)
            pdf.set_font("cjk", size=8)
        if len(bands) > 1:
            names = [str(headers[i]) if i < len(headers) else f"列{i + 1}"
                     for i in indices]
            label = f"列组 {band_no}/{len(bands)}：" + " / ".join(names)
            label_lines, _ = _pdf_wrap(pdf, label, avail, 2)
            for line in label_lines:
                write_cell(avail, 5, line, "L", next_line=True)
            pdf.ln(1)
        if header_layout:
            draw_layout(*header_layout, indices, widths, header_row=True)

    first_page = True
    for band_no, indices in enumerate(bands, 1):
        widths = widths_for(indices)
        header_layout = layout(headers, indices, widths, True) if headers else None
        add_page(indices, widths, header_layout, band_no, first=first_page)
        first_page = False
        alternate = False
        for row in rows:
            row_layout = layout(row, indices, widths)
            if pdf.get_y() + row_layout[1] > pdf.h - pdf.b_margin:
                add_page(indices, widths, header_layout, band_no)
            draw_layout(*row_layout, indices, widths, alternate=alternate)
            alternate = not alternate

    warnings = list(meta["warnings"])
    if truncated:
        warnings.append(
            f"源数据 {total_rows} 行，只输出前 {max_rows} 行；如需完整内容请调高行数上限。")
    if len(bands) > 1:
        warnings.append(
            f"源表有 {ncol} 列，A4 无法在单页保持可读；已拆成 {len(bands)} 个列组，"
            "后续列组重复首列用于对齐。")
    if clipped_cells:
        warnings.append(
            f"有 {clipped_cells} 个超长单元格为保证分页最多显示 10 行，末尾已用省略号标记。")
    if clipped_headers:
        warnings.append(
            f"有 {clipped_headers} 个超长表头最多显示 5 行，末尾已用省略号标记。")

    out.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out))
    return {"file": out.name, "path": str(out), "rows": len(rows),
            "cols": ncol, "truncated": truncated,
            "size": f"{out.stat().st_size / 1024:.0f} KB",
            "font": Path(font_path).name,
            "emoji_font": emoji_font,
            "sheet": meta["sheet"],
            "sheets": [s["name"] for s in meta["sheets"]],
            "total_rows": total_rows, "warnings": warnings,
            "clipped_cells": clipped_cells,
            "clipped_headers": clipped_headers,
            "column_bands": len(bands)}


def _fit(pdf, text: str, width: float, cache: dict | None = None) -> str:
    """超宽文本截断加省略号，避免串格。

    逐字符缩短会对每个单元格反复测量宽度，几万个格子能跑上两分钟。
    这里先按宽度比例估算该保留多少字符，再微调几次；
    并对「文本+列宽」组合做缓存——数据表里重复值很多（类别名、状态名），
    缓存命中率很高。
    """
    text = text.replace("\n", " ").replace("\r", " ")
    if not text:
        return ""
    key = (text, round(width, 2))
    if cache is not None and key in cache:
        return cache[key]
    limit = width - 2
    w = pdf.get_string_width(text)
    if w <= limit:
        out = text
    else:
        keep = max(1, int(len(text) * limit / w)) if w else 1
        out = text[:keep]
        # 估算偏长就收，偏短就放，正常一两次就收敛
        while keep > 1 and pdf.get_string_width(out + "…") > limit:
            keep -= 1
            out = text[:keep]
        while keep < len(text) and pdf.get_string_width(text[:keep + 1] + "…") <= limit:
            keep += 1
            out = text[:keep]
        out += "…"
    if cache is not None:
        cache[key] = out
    return out


# ══════════════════════════════════════════════════════════════
# Excel → Word
# ══════════════════════════════════════════════════════════════


def excel_to_word(xlsx: Path, out: Path, landscape: bool = True,
                  max_rows: int = 3000, title: str = "", sheet=None) -> dict:
    """Excel → Word（.docx），按内容分配列宽并提供可读的分页表格。"""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    headers, rows, meta = core.read_sheet_meta(xlsx, sheet)
    if not headers and not rows:
        raise ValueError("没读到任何内容。" + "；".join(meta["warnings"])
                         + f"\n工作表清单: {[s['name'] for s in meta['sheets']]}")
    total_rows = len(rows)
    truncated = total_rows > max_rows
    rows = rows[:max_rows]

    doc = Document()
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)

    word_font = ("PingFang SC" if sys.platform == "darwin" else
                 ("Microsoft YaHei" if sys.platform == "win32" else "Noto Sans CJK SC"))

    def style_run(run, size, bold=False, color=None):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = word_font
        run._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), word_font)
        if color is not None:
            run.font.color.rgb = color

    if title:
        h = doc.add_paragraph(title)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.runs[0]
        style_run(run, 14, bold=True, color=RGBColor(31, 55, 88))
        h.paragraph_format.space_after = Pt(8)

    ncol = max(len(headers), max((len(r) for r in rows), default=0)) or 1
    available_width = int(sec.page_width - sec.left_margin - sec.right_margin)
    weights = []
    for i in range(ncol):
        column_header = headers[i] if i < len(headers) else ""
        values = [column_header]
        values += [r[i] for r in rows[:200] if i < len(r)]
        weight = min(max(max((_content_weight(v, column_header) for v in values),
                             default=4), 4), 48)
        header_text = str(column_header).lower()
        if any(k in header_text for k in ("日期", "时间", "date", "time")):
            weight = max(weight, 12)
        if any(k in header_text for k in ("金额", "价格", "费用", "amount", "price")):
            weight = max(weight, 13)
        if any(k in header_text for k in ("百分", "比例")) or header_text.endswith("率"):
            weight = max(weight, 10)
        if any(k in header_text for k in ("电话", "手机", "联系方式")):
            weight = max(weight, 14)
        weights.append(weight)
    weights = [max(w, 5) for w in weights]

    # 与 PDF 一样，不能把几十列硬塞进一张 A4 页面。宽表按列组输出，
    # 后续列组重复首列，让编号/主键始终可用于横向对齐。
    max_band_cols = 12 if landscape else 7
    if ncol <= max_band_cols:
        bands = [list(range(ncol))]
    else:
        payload_size = max_band_cols - 1
        bands = [[0, *range(start, min(start + payload_size, ncol))]
                 for start in range(1, ncol, payload_size)]

    def widths_for(indices):
        selected = [weights[i] for i in indices]
        total_weight = sum(selected) or 1
        return [int(available_width * w / total_weight) for w in selected]

    def set_cell_width(cell, width):
        cell.width = width
        tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tc_w.set(qn("w:w"), str(max(1, int(width / 635))))
        tc_w.set(qn("w:type"), "dxa")

    def set_cell_margins(cell, top=55, start=75, bottom=55, end=75):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for edge, value in (("top", top), ("start", start),
                            ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def keep_row_together(row):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    def repeat_header(row):
        tr_pr = row._tr.get_or_add_trPr()
        flag = OxmlElement("w:tblHeader")
        flag.set(qn("w:val"), "true")
        tr_pr.append(flag)

    def shade(cell, color: str):
        el = OxmlElement("w:shd")
        el.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(el)

    for band_no, indices in enumerate(bands, 1):
        if band_no > 1:
            doc.add_page_break()
        if len(bands) > 1:
            names = [str(headers[i]) if i < len(headers) else f"列{i + 1}"
                     for i in indices]
            label = doc.add_paragraph(
                f"列组 {band_no}/{len(bands)}：" + " / ".join(names))
            label.paragraph_format.space_before = Pt(0)
            label.paragraph_format.space_after = Pt(5)
            label.paragraph_format.keep_with_next = True
            for run in label.runs:
                style_run(run, 9, bold=True, color=RGBColor(47, 84, 150))

        table = doc.add_table(rows=1 if headers else 0, cols=len(indices))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)

        word_widths = widths_for(indices)
        for local_i, width in enumerate(word_widths):
            table.columns[local_i].width = width
        body_size = 8.5 if len(indices) <= 8 else 7.5

        if headers:
            header_row = table.rows[0]
            repeat_header(header_row)
            keep_row_together(header_row)
            for local_i, original_i in enumerate(indices):
                txt = str(headers[original_i]) if original_i < len(headers) else ""
                cell = header_row.cells[local_i]
                cell.text = txt
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_width(cell, word_widths[local_i])
                set_cell_margins(cell, top=70, bottom=70)
                shade(cell, "2F5496")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        style_run(run, 9, bold=True,
                                  color=RGBColor(255, 255, 255))

        for ri, source_row in enumerate(rows):
            row = table.add_row()
            keep_row_together(row)
            for local_i, original_i in enumerate(indices):
                value = (source_row[original_i]
                         if original_i < len(source_row) else None)
                header = (str(headers[original_i])
                          if original_i < len(headers) else "")
                cell = row.cells[local_i]
                cell.text = _display_value(value, header)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_width(cell, word_widths[local_i])
                set_cell_margins(cell)
                if ri % 2:
                    shade(cell, "F5F6F8")
                for p in cell.paragraphs:
                    if isinstance(value, (int, float)) and not isinstance(value, bool):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif any(k in header for k in
                             ("日期", "时间", "状态", "比例", "百分")):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1
                    for run in p.runs:
                        style_run(run, body_size, color=RGBColor(32, 39, 51))

    warnings = list(meta["warnings"])
    if truncated:
        warnings.append(
            f"源数据 {total_rows} 行，只输出前 {max_rows} 行；如需完整内容请调高行数上限。")
    if len(bands) > 1:
        warnings.append(
            f"源表有 {ncol} 列，A4 无法在单页保持可读；已拆成 {len(bands)} 个列组，"
            "后续列组重复首列用于对齐。")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return {"file": out.name, "path": str(out), "rows": len(rows),
            "cols": ncol, "truncated": truncated,
            "size": f"{out.stat().st_size / 1024:.0f} KB",
            "sheet": meta["sheet"],
            "sheets": [s["name"] for s in meta["sheets"]],
            "total_rows": total_rows, "warnings": warnings,
            "column_bands": len(bands)}


# ══════════════════════════════════════════════════════════════
# 图片 → Excel（macOS Vision / Windows OCR）
# ══════════════════════════════════════════════════════════════

OCR_SWIFT = r'''
import Foundation
import Vision
import AppKit

let args = CommandLine.arguments
guard args.count > 1, let img = NSImage(contentsOfFile: args[1]),
      let tiff = img.tiffRepresentation,
      let bmp = NSBitmapImageRep(data: tiff),
      let cg = bmp.cgImage else {
    FileHandle.standardError.write("无法读取图片\n".data(using: .utf8)!)
    exit(2)
}
let req = VNRecognizeTextRequest()
req.recognitionLevel = .accurate
req.recognitionLanguages = ["zh-Hans", "en-US"]
req.usesLanguageCorrection = false
let handler = VNImageRequestHandler(cgImage: cg, options: [:])
do { try handler.perform([req]) } catch {
    FileHandle.standardError.write("OCR 失败: \(error)\n".data(using: .utf8)!)
    exit(3)
}
var out: [[String: Any]] = []
for ob in (req.results ?? []) {
    guard let top = ob.topCandidates(1).first else { continue }
    let b = ob.boundingBox            // 归一化坐标，原点在左下
    out.append(["text": top.string, "conf": top.confidence,
                "x": b.minX, "y": 1 - b.maxY, "w": b.width, "h": b.height])
}
print(String(data: try! JSONSerialization.data(withJSONObject: out),
             encoding: .utf8)!)
'''


OCR_POWERSHELL = r'''
param([Parameter(Mandatory=$true)][string]$ImagePath)
$ErrorActionPreference = 'Stop'
Add-Type -AssemblyName System.Runtime.WindowsRuntime
$null = [Windows.Storage.StorageFile, Windows.Storage, ContentType=WindowsRuntime]
$null = [Windows.Storage.FileAccessMode, Windows.Storage, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Graphics.Imaging, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.SoftwareBitmap, Windows.Graphics.Imaging, ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrResult, Windows.Foundation, ContentType=WindowsRuntime]

$script:asTask = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {
  $_.Name -eq 'AsTask' -and $_.IsGenericMethod -and
  $_.GetGenericArguments().Count -eq 1 -and $_.GetParameters().Count -eq 1
})[0]
function Await-WinRt($Operation, [Type]$ResultType) {
  $task = $script:asTask.MakeGenericMethod($ResultType).Invoke($null, @($Operation))
  $task.Wait()
  return $task.Result
}

$file = Await-WinRt ([Windows.Storage.StorageFile]::GetFileFromPathAsync($ImagePath)) ([Windows.Storage.StorageFile])
$stream = Await-WinRt ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
$decoder = Await-WinRt ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
$bitmap = Await-WinRt ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
$engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
if ($null -eq $engine) { throw 'Windows OCR language pack is unavailable' }
$result = Await-WinRt ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])

$blocks = @()
foreach ($line in $result.Lines) {
  $words = @($line.Words | Sort-Object { $_.BoundingRect.X })
  $group = @()
  $right = 0.0
  $height = 0.0
  foreach ($word in $words) {
    $rect = $word.BoundingRect
    $gap = [double]$rect.X - $right
    # 中文 OCR 偶尔把同一表头拆成单字/短词；允许约 1.35 个字高的词间距，
    # 同时仍远小于常见单元格之间的留白。
    $joinGap = [Math]::Max(6.0, [Math]::Min([double]$height, [double]$rect.Height) * 1.35)
    if ($group.Count -gt 0 -and $gap -gt $joinGap) {
      $first = $group[0].BoundingRect
      $last = $group[-1].BoundingRect
      $top = ($group | ForEach-Object { $_.BoundingRect.Y } | Measure-Object -Minimum).Minimum
      $bottom = ($group | ForEach-Object { $_.BoundingRect.Y + $_.BoundingRect.Height } | Measure-Object -Maximum).Maximum
      $blocks += [ordered]@{
        text = (($group | ForEach-Object { $_.Text }) -join ' ')
        x = [double]$first.X / [double]$bitmap.PixelWidth
        y = [double]$top / [double]$bitmap.PixelHeight
        w = ([double]$last.X + [double]$last.Width - [double]$first.X) / [double]$bitmap.PixelWidth
        h = ([double]$bottom - [double]$top) / [double]$bitmap.PixelHeight
      }
      $group = @()
    }
    $group += $word
    $right = [double]$rect.X + [double]$rect.Width
    $height = [double]$rect.Height
  }
  if ($group.Count -gt 0) {
    $first = $group[0].BoundingRect
    $last = $group[-1].BoundingRect
    $top = ($group | ForEach-Object { $_.BoundingRect.Y } | Measure-Object -Minimum).Minimum
    $bottom = ($group | ForEach-Object { $_.BoundingRect.Y + $_.BoundingRect.Height } | Measure-Object -Maximum).Maximum
    $blocks += [ordered]@{
      text = (($group | ForEach-Object { $_.Text }) -join ' ')
      x = [double]$first.X / [double]$bitmap.PixelWidth
      y = [double]$top / [double]$bitmap.PixelHeight
      w = ([double]$last.X + [double]$last.Width - [double]$first.X) / [double]$bitmap.PixelWidth
      h = ([double]$bottom - [double]$top) / [double]$bitmap.PixelHeight
    }
  }
}
ConvertTo-Json -InputObject @($blocks) -Compress
'''


def _ensure_ocr_binary() -> Path:
    """确保 OCR 二进制存在：优先用随包预编译的，否则就地编译并缓存"""
    bundled = core.res_path("vendor/ocr")
    if bundled.exists():
        return bundled
    import hashlib
    tag = hashlib.sha256(OCR_SWIFT.encode()).hexdigest()[:16]
    cache_dir = core.CACHE_DIR / f"ocr_{tag}"
    binary = cache_dir / "ocr"
    if binary.exists():
        return binary
    if not _has_swiftc():
        raise RuntimeError(
            "图片识别需要 macOS 的 Swift 编译器（Xcode Command Line Tools）。\n"
            "请在终端执行 xcode-select --install 后重试，"
            "或改用「文本转 Excel」手工粘贴表格内容。")
    cache_dir.mkdir(parents=True, exist_ok=True)
    src = cache_dir / "ocr.swift"
    src.write_text(OCR_SWIFT, encoding="utf-8")
    r = subprocess.run(["swiftc", "-O", str(src), "-o", str(binary)],
                       capture_output=True, text=True, timeout=300,
                       env=core.clean_subprocess_env())
    if r.returncode != 0 or not binary.exists():
        raise RuntimeError("编译 OCR 程序失败:\n" + (r.stderr or "")[:600])
    return binary


def _has_swiftc() -> bool:
    from shutil import which
    return which("swiftc") is not None


def _windows_ocr_blocks(image: Path) -> list[dict]:
    """调用 Windows.Media.Ocr；脚本写入带 BOM 的缓存文件以兼容 PowerShell 5。"""
    import hashlib

    tag = hashlib.sha256(OCR_POWERSHELL.encode()).hexdigest()[:16]
    cache_dir = core.CACHE_DIR / f"win_ocr_{tag}"
    cache_dir.mkdir(parents=True, exist_ok=True)
    script = cache_dir / "ocr.ps1"
    if not script.exists():
        script.write_text(OCR_POWERSHELL, encoding="utf-8-sig")
    r = subprocess.run(
        ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
         "-File", str(script), "-ImagePath", str(image.resolve())],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180,
        env=core.clean_subprocess_env(),
    )
    if r.returncode != 0:
        detail = (r.stderr or r.stdout or "未知错误")[:500]
        raise RuntimeError(
            "Windows 图片识别失败。请确认系统已安装中文 OCR 语言包。\n" + detail)
    line = next((item for item in r.stdout.splitlines()
                 if item.lstrip().startswith("[")), None)
    if line is None:
        raise RuntimeError("图片识别没有返回结果，可能图里没有可识别的文字")
    return json.loads(line)


def ocr_blocks(image: Path) -> list[dict]:
    """调用当前系统的原生 OCR，返回带归一化坐标的文字块。"""
    if not image.exists():
        raise FileNotFoundError(f"图片不存在: {image}")
    if sys.platform == "win32":
        return _windows_ocr_blocks(image)
    if sys.platform != "darwin":
        raise RuntimeError("图片识别当前支持 Windows 10/11 和 macOS")
    binary = _ensure_ocr_binary()
    r = subprocess.run([str(binary), str(image)], capture_output=True,
                       text=True, timeout=180, env=core.clean_subprocess_env())
    if r.returncode != 0:
        raise RuntimeError("图片识别失败: " + (r.stderr or "未知错误")[:400])
    line = next((l for l in r.stdout.splitlines() if l.startswith("[")), None)
    if line is None:
        raise RuntimeError("图片识别没有返回结果，可能图里没有可识别的文字")
    return json.loads(line)


def blocks_to_rows(blocks: list[dict], row_tol: float = 0.6,
                   col_gap: float = 0.025, min_fill: float = 0.15) -> list[list[str]]:
    """把带坐标的文字块还原成表格二维数组。

    行：按 y 坐标聚类，容差取字高的一定比例（同一行的文字 y 几乎相同）。
    列：把所有块的 x 起点聚类成列锚点，再把每个块归到最近的锚点。
    最后剔除填充率过低的列——截图里的水印、翻页控件会各自占一个稀疏列，
    不剔掉的话两列表格会变成七列。
    """
    if not blocks:
        return []

    def clean_text(value) -> str:
        text = str(value or "").strip()
        # Windows OCR 的 Word 集合常输出“交 易 日 期”。中文字符间的空格
        # 不是业务内容，去掉后表头与字段名才能稳定匹配；英文词间空格保留。
        return re.sub(r"(?<=[\u3400-\u9fff])\s+(?=[\u3400-\u9fff])", "", text)

    items = sorted(blocks, key=lambda b: (b["y"], b["x"]))

    # ── 聚行 ──
    lines: list[list[dict]] = []
    for b in items:
        placed = False
        for ln in lines:
            ref = ln[0]
            tol = max(ref["h"], b["h"]) * row_tol
            if abs((b["y"] + b["h"] / 2) - (ref["y"] + ref["h"] / 2)) <= tol:
                ln.append(b)
                placed = True
                break
        if not placed:
            lines.append([b])
    for ln in lines:
        ln.sort(key=lambda b: b["x"])
    lines.sort(key=lambda ln: min(b["y"] for b in ln))

    # ── 聚列：所有 x 起点归并成锚点 ──
    xs = sorted(b["x"] for b in blocks)
    anchors: list[float] = []
    for x in xs:
        if not anchors or x - anchors[-1] > col_gap:
            anchors.append(x)
        else:
            anchors[-1] = (anchors[-1] + x) / 2
    if not anchors:
        return [[clean_text(b["text"]) for b in ln] for ln in lines]

    rows: list[list[str]] = []
    for ln in lines:
        cells = [""] * len(anchors)
        for b in ln:
            ci = min(range(len(anchors)), key=lambda i: abs(anchors[i] - b["x"]))
            text = clean_text(b["text"])
            cells[ci] = (cells[ci] + " " + text).strip() if cells[ci] else text
        rows.append(cells)

    if not rows:
        return []
    # 按填充率保留列：至少留 1 列，避免全被剔掉
    total = len(rows)
    fill_rate = [sum(1 for r in rows if r[i].strip()) / total
                 for i in range(len(anchors))]
    keep = [i for i, fr in enumerate(fill_rate) if fr >= min_fill]
    if not keep:
        keep = [max(range(len(fill_rate)), key=lambda i: fill_rate[i])]
    rows = [[r[i] for i in keep] for r in rows]
    # 去掉剔列后变成全空的行
    return [r for r in rows if any(c.strip() for c in r)]


def image_to_rows(image: Path, row_tol: float = 0.6, col_gap: float = 0.025,
                  min_fill: float = 0.15) -> dict:
    """图片 → 表格二维数组 + 识别质量信息"""
    blocks = ocr_blocks(image)
    if not blocks:
        raise ValueError("图片中没有识别到文字；请换用清晰、正向且裁剪到表格区域的截图")
    rows = blocks_to_rows(blocks, row_tol, col_gap, min_fill)
    confidence_available = bool(blocks) and all(
        isinstance(b.get("conf"), (int, float)) for b in blocks)
    confs = [float(b["conf"]) for b in blocks] if confidence_available else []
    low = ([b["text"] for b in blocks if float(b["conf"]) < 0.5]
           if confidence_available else [])
    warnings = []
    if not confidence_available and blocks:
        warnings.append("当前系统 OCR 不提供逐块置信度，生成 Excel 前请逐格人工核对")
    if low:
        warnings.append("部分文字识别置信度偏低，请重点核对相似字符 f/t、l/1、0/O")
    if len(rows) <= 1:
        warnings.append("只识别到 1 行，无法可靠判断表格行结构，请核对是否漏行或把表头拆列")
    if rows and len(rows[0]) <= 1:
        warnings.append("只识别到 1 列，图片可能过小、模糊或有旋转，请更换清晰正向截图")
    return {
        "rows": rows,
        "block_count": len(blocks),
        "row_count": len(rows),
        "col_count": len(rows[0]) if rows else 0,
        "avg_conf": round(sum(confs) / len(confs), 3) if confs else None,
        "confidence_available": confidence_available,
        "low_conf_texts": low[:20],
        "warning": "；".join(warnings),
    }
