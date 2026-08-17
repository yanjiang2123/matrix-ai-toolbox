import csv
import tempfile
import unittest
import zipfile
from datetime import date, datetime
from pathlib import Path
from unittest.mock import patch

import openpyxl

import convert
import excel_diff
import matrix_core
import sql_tools


class DifferenceArchiveTests(unittest.TestCase):
    def test_difference_moves_through_fixed_and_reappeared_states(self):
        diff = {"pk": "1", "col": "amount", "a": "10", "b": "11",
                "kind": "字段值差异"}
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "archive.xlsx"

            first = matrix_core.accumulate_diffs(path, [diff], "第一轮", ["1"])
            self.assertEqual((1, 0), (first["added"], first["fixed"]))
            self.assertEqual("未修复", self._status(path))

            clean = matrix_core.accumulate_diffs(path, [], "第二轮", ["1"])
            self.assertEqual(1, clean["fixed"])
            self.assertEqual("已修复", self._status(path))

            again = matrix_core.accumulate_diffs(path, [diff], "第三轮", ["1"])
            self.assertEqual(1, again["updated"])
            self.assertEqual("又出现", self._status(path))

            fixed_again = matrix_core.accumulate_diffs(path, [], "第四轮", ["1"])
            self.assertEqual(1, fixed_again["fixed"])
            self.assertEqual("已修复", self._status(path))

    @staticmethod
    def _status(path):
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        try:
            return wb.active.cell(2, 8).value
        finally:
            wb.close()


class ExportSafetyTests(unittest.TestCase):
    def test_xlsx_and_csv_escape_formula_like_values(self):
        with tempfile.TemporaryDirectory() as tmp:
            xlsx_path = Path(tmp) / "safe.xlsx"
            csv_path = Path(tmp) / "safe.csv"
            matrix_core.write_xlsx(
                xlsx_path, {"数据": (["=危险表头"], [["=1+1"], [" +cmd"]])}
            )
            matrix_core.write_csv(csv_path, ["=危险表头"], [["=1+1"], [" +cmd"]])

            wb = openpyxl.load_workbook(xlsx_path, data_only=False)
            try:
                ws = wb["数据"]
                self.assertEqual("'=危险表头", ws["A1"].value)
                self.assertEqual("'=1+1", ws["A2"].value)
                self.assertEqual("s", ws["A2"].data_type)
                self.assertEqual("' +cmd", ws["A3"].value)
            finally:
                wb.close()

            with csv_path.open(encoding="utf-8-sig", newline="") as handle:
                rows = list(csv.reader(handle))
            self.assertEqual([["'=危险表头"], ["'=1+1"], ["' +cmd"]], rows)


class TextExcelConversionTests(unittest.TestCase):
    def test_high_precision_decimal_text_is_not_rounded_by_excel(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "precision.xlsx"
            matrix_core.write_xlsx(path, {"数据": (
                ["金额", "超大值", "比例", "科学小数", "科学大数"],
                [["-123456789012345.6789012345",
                  "99999999999999999999.9999999999",
                  "12.34567890123456789%", "-9.87654321E-25",
                  "4.93827156E+18"]],
            )})
            wb = openpyxl.load_workbook(path, data_only=False)
            try:
                ws = wb["数据"]
                self.assertEqual("-123456789012345.6789012345", ws["A2"].value)
                self.assertEqual("99999999999999999999.9999999999", ws["B2"].value)
                self.assertEqual("12.34567890123456789%", ws["C2"].value)
                self.assertEqual("@", ws["A2"].number_format)
                self.assertEqual("@", ws["B2"].number_format)
                self.assertEqual("@", ws["C2"].number_format)
                self.assertAlmostEqual(-9.87654321e-25, ws["D2"].value)
                self.assertAlmostEqual(4.93827156e18, ws["E2"].value)
                self.assertEqual("0.00000000E+00", ws["D2"].number_format)
                self.assertEqual("0.00000000E+00", ws["E2"].number_format)
            finally:
                wb.close()

    def test_single_uncached_formula_is_not_silently_lost(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "formula.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["id", "计算结果"])
            ws.append([1, "=1+1"])
            wb.save(path)
            wb.close()

            headers, rows, meta = matrix_core.read_sheet_meta(path)
            self.assertEqual(["id", "计算结果"], headers)
            self.assertEqual("=1+1", rows[0][1])
            self.assertEqual(1, meta["formula_cells"])
            self.assertTrue(any("公式原文" in w for w in meta["warnings"]))

    def test_quoted_delimiter_and_internal_blank_row_are_preserved(self):
        rows = matrix_core.text_to_rows(
            '\ufeff编号,备注,金额\n1,"包含,逗号",12.50\n\n2,"普通文本",8', ","
        )
        self.assertEqual(["编号", "备注", "金额"], rows[0])
        self.assertEqual(["1", "包含,逗号", "12.50"], rows[1])
        self.assertEqual(["", "", ""], rows[2])
        self.assertEqual(["2", "普通文本", "8"], rows[3])

    def test_xlsx_has_business_types_and_readable_formatting(self):
        long_note = "用于检查自动换行与行高的中文长文本。" * 12
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "formatted.xlsx"
            matrix_core.write_xlsx(path, {"数据": (
                ["编号", "日期", "金额", "比例", "备注", "命令"],
                [
                    ["001", "2026-08-17", "1234.50", "12.5%", long_note, "=1+1"],
                    ["002", "", "-45.67", "0%", "短文本", " +cmd"],
                ],
            )})

            wb = openpyxl.load_workbook(path, data_only=False)
            try:
                ws = wb["数据"]
                self.assertEqual("001", ws["A2"].value)
                self.assertEqual("s", ws["A2"].data_type)
                self.assertEqual("@", ws["A2"].number_format)
                self.assertIsInstance(ws["B2"].value, date)
                self.assertEqual("yyyy-mm-dd", ws["B2"].number_format)
                self.assertEqual(1234.5, ws["C2"].value)
                self.assertEqual("#,##0.00", ws["C2"].number_format)
                self.assertEqual(-45.67, ws["C3"].value)
                self.assertEqual(0.125, ws["D2"].value)
                self.assertEqual("0.0%", ws["D2"].number_format)
                self.assertEqual("'=1+1", ws["F2"].value)
                self.assertEqual("' +cmd", ws["F3"].value)
                self.assertEqual("微软雅黑", ws["E2"].font.name)
                self.assertTrue(ws["E2"].alignment.wrap_text)
                self.assertEqual("thin", ws["E2"].border.left.style)
                self.assertGreater(ws.row_dimensions[2].height, 20)
                self.assertLessEqual(ws.column_dimensions["E"].width, 42)
                self.assertEqual("A2", ws.freeze_panes)
                self.assertEqual("A1:F3", ws.auto_filter.ref)
                self.assertEqual(1, ws.page_setup.fitToWidth)
                self.assertIn("A1:F3", str(ws.print_area).replace("$", ""))
            finally:
                wb.close()


class ReportAndSqlTests(unittest.TestCase):
    def test_same_display_names_keep_both_side_sheets(self):
        result = excel_diff.compare_tables(
            ["id", "amount"], [["A", 1]],
            ["id", "amount"], [["B", 2]],
            ["id"],
        )
        report = excel_diff.build_report(result, "same.xlsx", "same.xlsx")
        self.assertIn("仅A有", report)
        self.assertIn("仅B有", report)
        self.assertEqual("A", report["仅A有"][1][0][0])
        self.assertEqual("B", report["仅B有"][1][0][0])

    def test_insert_quotes_each_table_identifier(self):
        sql = convert._wrap_insert("sales.order", "`id`", ["(1)"])
        self.assertTrue(sql.startswith("INSERT INTO `sales`.`order`"))

    def test_ddl_special_identifier_and_decimal_scale_are_preserved(self):
        ddl = "CREATE TABLE demo (`金额（元）` decimal(18,2), `业务日期` date)"
        parsed = convert.parse_ddl(ddl)
        self.assertEqual(["金额（元）", "业务日期"],
                         [col["name"] for col in parsed["columns"]])
        self.assertEqual(2, parsed["columns"][0]["scale"])
        self.assertEqual("6172.80", convert._sql_literal(
            6172.799999999999, parsed["columns"][0]))
        self.assertEqual("12.00", convert._sql_literal(
            12.0, parsed["columns"][0]))
        self.assertEqual("'2026-08-01'", convert._sql_literal(
            datetime(2026, 8, 1), parsed["columns"][1]))

        wide = convert.parse_ddl(
            "CREATE TABLE demo (`amount` decimal(38,18))"
        )["columns"][0]
        self.assertEqual("4938271605000000000.000000000000000000",
                         convert._sql_literal(4938271605000000000, wide))

    def test_insert_preserves_literal_null_in_text_columns(self):
        text_col = {"type": "varchar", "quoted": True, "scale": None}
        numeric_col = {"type": "decimal", "quoted": False, "scale": 2}
        date_col = {"type": "date", "quoted": True, "scale": None}
        self.assertEqual("'NULL'", convert._sql_literal("NULL", text_col))
        self.assertEqual("NULL", convert._sql_literal("NULL", numeric_col))
        self.assertEqual("NULL", convert._sql_literal("NULL", date_col))
        self.assertEqual("NULL", convert._sql_literal("\\N", text_col))

    def test_insert_rejects_invalid_numeric_values_instead_of_quoting_them(self):
        numeric_col = {"name": "amount", "type": "decimal",
                       "quoted": False, "scale": 2}
        self.assertEqual("1234.50", convert._sql_literal("1,234.50", numeric_col))
        for value in ("not-a-number", "NaN", "Infinity"):
            with self.subTest(value=value):
                with self.assertRaisesRegex(ValueError, "非数值内容"):
                    convert._sql_literal(value, numeric_col)

    def test_insert_rejects_decimal_values_outside_declared_precision(self):
        col = convert.parse_ddl(
            "CREATE TABLE demo (`amount` DECIMAL(5,2))"
        )["columns"][0]
        self.assertEqual(5, col["precision"])
        self.assertEqual("999.99", convert._sql_literal("999.99", col))
        with self.assertRaisesRegex(ValueError, r"超出 DECIMAL\(5,2\) 范围"):
            convert._sql_literal("1000", col)
        with self.assertRaisesRegex(ValueError, r"超出 DECIMAL\(5,2\) 范围"):
            convert._sql_literal("999.995", col)

    def test_insert_rejects_excel_floats_that_cannot_keep_long_ids_exact(self):
        col = {"name": "business_id", "type": "bigint",
               "quoted": False, "precision": None, "scale": None}
        with self.assertRaisesRegex(ValueError, "超过 15 位"):
            convert._sql_literal(1.2345678901234568e18, col)
        self.assertEqual("1234567890123456789",
                         convert._sql_literal(1234567890123456789, col))

    def test_word_has_readable_header_pagination_and_display_formats(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.xlsx"
            output = Path(tmp) / "result.docx"
            matrix_core.write_xlsx(source, {"数据": (
                ["日期", "金额", "完成率", "状态", "长说明"],
                [[date(2026, 8, 1), 1280.5, 0.125, True, "中文说明" * 30]],
            )})
            result = convert.excel_to_word(source, output, title="业务验收")
            self.assertEqual(1, result["rows"])
            self.assertTrue(output.exists())

            from docx import Document
            doc = Document(output)
            table = doc.tables[0]
            self.assertEqual("FFFFFF", str(
                table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb))
            self.assertEqual("2026-08-01", table.rows[1].cells[0].text)
            self.assertEqual("1,280.50", table.rows[1].cells[1].text)
            self.assertEqual("12.50%", table.rows[1].cells[2].text)
            self.assertEqual("是", table.rows[1].cells[3].text)
            with zipfile.ZipFile(output) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("<w:tblHeader", xml)
            self.assertIn("<w:cantSplit", xml)
            self.assertIn('<w:tblLayout w:type="fixed"', xml)
            self.assertIn("<w:tcMar", xml)

    def test_word_splits_wide_tables_into_readable_column_bands(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "wide.xlsx"
            output = Path(tmp) / "wide.docx"
            headers = ["业务主键", *[f"字段{i}" for i in range(1, 14)]]
            matrix_core.write_xlsx(source, {"数据": (
                headers, [["001", *[f"值{i}" for i in range(1, 14)]]],
            )})
            result = convert.excel_to_word(source, output, landscape=True)
            self.assertEqual(2, result["column_bands"])
            self.assertTrue(any("列组" in w for w in result["warnings"]))

            from docx import Document
            doc = Document(output)
            self.assertEqual(2, len(doc.tables))
            self.assertTrue(all(len(table.columns) <= 12 for table in doc.tables))
            self.assertEqual("业务主键", doc.tables[0].rows[0].cells[0].text)
            self.assertEqual("业务主键", doc.tables[1].rows[0].cells[0].text)
            self.assertEqual("001", doc.tables[1].rows[1].cells[0].text)

    def test_pdf_can_use_available_local_font_and_wrap_long_text(self):
        if convert._pick_font() is None:
            self.skipTest("本机没有可供 PDF 使用的字体")
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.xlsx"
            output = Path(tmp) / "result.pdf"
            matrix_core.write_xlsx(source, {"数据": (
                ["日期", "金额", "说明"],
                [[date(2026, 8, 1), 1280.5, "需要自动换行的业务说明" * 8]],
            )})
            result = convert.excel_to_pdf(source, output, title="业务验收")
            self.assertEqual(1, result["rows"])
            self.assertEqual(0, result["clipped_cells"])
            self.assertEqual(0, result["clipped_headers"])
            self.assertGreater(output.stat().st_size, 1000)

    def test_pdf_splits_wide_tables_into_readable_column_bands(self):
        if convert._pick_font() is None:
            self.skipTest("本机没有可供 PDF 使用的字体")
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "wide.xlsx"
            output = Path(tmp) / "wide.pdf"
            headers = ["业务主键", *[f"字段{i}" for i in range(1, 14)]]
            matrix_core.write_xlsx(source, {"数据": (
                headers, [["001", *[f"值{i}" for i in range(1, 14)]]],
            )})
            result = convert.excel_to_pdf(source, output, landscape=True)
            self.assertEqual(2, result["column_bands"])
            self.assertTrue(any("列组" in w for w in result["warnings"]))
            self.assertGreater(output.stat().st_size, 1000)

    def test_native_ocr_without_confidence_is_labeled_for_manual_review(self):
        blocks = [
            {"text": "编号", "x": 0.1, "y": 0.1, "w": 0.1, "h": 0.05},
            {"text": "001", "x": 0.1, "y": 0.2, "w": 0.1, "h": 0.05},
        ]
        with patch.object(convert, "ocr_blocks", return_value=blocks):
            result = convert.image_to_rows(Path("unused.png"))
        self.assertFalse(result["confidence_available"])
        self.assertIsNone(result["avg_conf"])
        self.assertIn("人工核对", result["warning"])
        self.assertIn("只识别到 1 列", result["warning"])

    def test_ocr_removes_artificial_spaces_between_chinese_characters(self):
        rows = convert.blocks_to_rows([
            {"text": "交 易 日 期", "x": 0.1, "y": 0.1,
             "w": 0.2, "h": 0.05},
            {"text": "Order Date", "x": 0.5, "y": 0.1,
             "w": 0.2, "h": 0.05},
        ])
        self.assertEqual([["交易日期", "Order Date"]], rows)

    def test_empty_ocr_result_is_an_explicit_failure(self):
        with patch.object(convert, "ocr_blocks", return_value=[]):
            with self.assertRaisesRegex(ValueError, "没有识别到文字"):
                convert.image_to_rows(Path("unused.png"))

    def test_sample_comparison_cannot_claim_full_equivalence(self):
        result = sql_tools.compare_details(
            ["id", "value"], [["1", "same"]],
            ["id", "value"], [["1", "same"]],
            ["id"],
        )
        sql_tools.qualify_sample_result(result, 500)
        self.assertEqual("warn", result["verdict"]["level"])
        self.assertTrue(result["verdict"]["sample_only"])
        self.assertIn("抽样结果", result["verdict"]["text"])
        self.assertIn("没有稳定 ORDER BY", result["verdict"]["next"])
        self.assertIn("全量分片", result["verdict"]["next"])

    def test_comparison_returns_all_checked_primary_keys(self):
        result = sql_tools.compare_details(
            ["id", "value"], [["1", "same"], ["2", "A"]],
            ["id", "value"], [["1", "same"], ["2", "B"]],
            ["id"],
        )
        self.assertEqual(["1", "2"], result["pks_in_scope"])

    def test_sql_detail_duplicate_groups_match_by_content_not_row_order(self):
        result = sql_tools.compare_details(
            ["id", "value"], [["1", "x"], ["1", "y"]],
            ["id", "value"], [["1", "y"], ["1", "x"]],
            ["id"],
        )
        self.assertEqual(2, result["stats"]["matched"])
        self.assertEqual(0, result["stats"]["diff_cells"])
        self.assertEqual(0, result["stats"]["only_a"])
        self.assertEqual(0, result["stats"]["only_b"])
        self.assertEqual("warn", result["verdict"]["level"])
        self.assertEqual(1, result["stats"]["duplicate_keys_a"])
        self.assertIn("主键不唯一", result["verdict"]["text"])

    def test_sql_detail_extra_duplicate_rows_are_not_silently_lost(self):
        cases = (
            ([["1", "x"], ["1", "x"]], [["1", "x"]], "only_a"),
            ([["1", "x"]], [["1", "x"], ["1", "x"]], "only_b"),
        )
        for rows_a, rows_b, side in cases:
            with self.subTest(side=side):
                result = sql_tools.compare_details(
                    ["id", "value"], rows_a,
                    ["id", "value"], rows_b,
                    ["id"],
                )
                self.assertEqual(1, result["stats"][side])
                self.assertEqual("diff", result["verdict"]["level"])

    def test_sql_detail_distinguishes_null_empty_and_literal_null(self):
        for value_b in ("", "NULL", "NONE"):
            with self.subTest(value_b=value_b):
                result = sql_tools.compare_details(
                    ["id", "value"], [["1", None]],
                    ["id", "value"], [["1", value_b]],
                    ["id"],
                )
                self.assertEqual(1, result["stats"]["diff_cells"])
                self.assertEqual(1, result["stats"]["null_flip"])
                self.assertEqual("（数据库 NULL）", result["diffs"][0]["a"])
                self.assertEqual("value", result["diffs"][0]["col"])

    def test_zero_padded_primary_keys_do_not_collapse_into_numbers(self):
        result = sql_tools.compare_details(
            ["id", "value"], [["001", "A"]],
            ["id", "value"], [["1", "A"]],
            ["id"],
        )
        self.assertEqual(1, result["stats"]["only_a"])
        self.assertEqual(1, result["stats"]["only_b"])
        self.assertEqual(["001"], result["only_a"])
        self.assertEqual(["1"], result["only_b"])

    def test_zero_padded_business_values_do_not_collapse_into_numbers(self):
        result = sql_tools.compare_details(
            ["id", "code"], [["1", "001"]],
            ["id", "code"], [["1", "1"]],
            ["id"],
        )
        self.assertEqual(1, result["stats"]["diff_cells"])
        self.assertEqual("code", result["diffs"][0]["col"])

    def test_sql_detail_does_not_claim_success_when_value_columns_differ(self):
        result = sql_tools.compare_details(
            ["id", "old_amount"], [["1", "10"]],
            ["id", "new_amount"], [["1", "10"]],
            ["id"],
        )
        self.assertEqual("diff", result["verdict"]["level"])
        self.assertEqual(["old_amount"], result["stats"]["only_cols_a"])
        self.assertEqual(["new_amount"], result["stats"]["only_cols_b"])
        self.assertEqual(0, result["stats"]["cmp_cols"])
        self.assertIn("字段结构不一致", result["verdict"]["text"])

    def test_sql_detail_warns_when_only_primary_keys_are_comparable(self):
        result = sql_tools.compare_details(
            ["id"], [["1"]], ["id"], [["1"]], ["id"])
        self.assertEqual("warn", result["verdict"]["level"])
        self.assertEqual(0, result["stats"]["cmp_cols"])
        self.assertIn("未比较任何非主键", result["verdict"]["text"])

    def test_sql_detail_does_not_treat_two_empty_results_as_proof(self):
        result = sql_tools.compare_details(
            ["id", "amount"], [], ["id", "amount"], [], ["id"])
        self.assertEqual("warn", result["verdict"]["level"])
        self.assertIn("0 行", result["verdict"]["text"])
        self.assertIn("时间窗", result["verdict"]["next"])

    def test_truncated_differences_are_excluded_from_archive_scope(self):
        value_columns = [f"c{i}" for i in range(2001)]
        headers = ["id", *value_columns]
        result = sql_tools.compare_details(
            headers, [["1", *(["A"] * len(value_columns))]],
            headers, [["1", *(["B"] * len(value_columns))]],
            ["id"],
        )
        self.assertTrue(result["archive_scope_truncated"])
        self.assertNotIn("1", result["pks_in_scope"])

    def test_duplicate_primary_key_headers_are_rejected(self):
        with self.assertRaisesRegex(ValueError, "主键字段在结果集中重名"):
            sql_tools.compare_details(
                ["id", "id", "amount"], [[1, 100, 10]],
                ["id", "id", "amount"], [[2, 100, 10]],
                ["id"],
            )

    def test_file_compare_rejects_duplicate_headers(self):
        with self.assertRaisesRegex(ValueError, "文件表头必须唯一"):
            excel_diff.compare_tables(
                ["id", "id", "value"], [[1, 100, "same"]],
                ["id", "id", "value"], [[2, 100, "same"]], ["id"])

    def test_composite_primary_key_text_cannot_collide_on_plus_sign(self):
        result = sql_tools.compare_details(
            ["k1", "k2", "value"], [["a+b", "c", "A"], ["a", "b+c", "A"]],
            ["k1", "k2", "value"], [["a+b", "c", "B"], ["a", "b+c", "B"]],
            ["k1", "k2"],
        )
        keys = [d["key"] for d in result["diffs"]]
        self.assertEqual(2, len(set(keys)))
        self.assertIn('["a+b","c"]', keys)
        self.assertIn('["a","b+c"]', keys)

    def test_one_side_empty_composite_keys_keep_component_boundaries(self):
        result = sql_tools.compare_details(
            ["k1", "k2", "value"],
            [["a+b", "c", "A"], ["a", "b+c", "A"]],
            [], [], ["k1", "k2"])
        self.assertEqual(2, len(result["pks_in_scope"]))
        self.assertIn('["a+b","c"]', result["only_a"])
        self.assertIn('["a","b+c"]', result["only_a"])

    def test_one_side_empty_rejects_duplicate_primary_key_headers(self):
        with self.assertRaisesRegex(ValueError, "主键字段重名"):
            sql_tools.compare_details(
                ["id", "id", "value"], [[1, 100, "A"]],
                [], [], ["id"])

    def test_incomplete_full_scan_cannot_close_archive_or_claim_success(self):
        plan = {"consistent": False, "planned_rows": 9, "total_rows": 10,
                "still_over": [{"start": "2026-08-01"}]}
        fetched = {"fetched": 8, "mismatch": [{"expect": 9, "got": 8}]}
        quality = sql_tools.assess_full_scan(plan, fetched, "A")
        self.assertFalse(quality["complete"])
        result = sql_tools.compare_details(
            ["id", "value"], [["1", "same"]],
            ["id", "value"], [["1", "same"]], ["id"])
        sql_tools.qualify_incomplete_scan(result, quality["issues"])
        self.assertEqual("warn", result["verdict"]["level"])
        self.assertFalse(result["verdict"]["scan_complete"])
        self.assertEqual([], result["pks_in_scope"])
        self.assertTrue(result["archive_scope_truncated"])
        self.assertIn("全量分片未完整覆盖", result["logs"][0]["reason"])

    def test_one_table_qualified_key_is_normalized_before_fetch(self):
        columns = [
            {"name": "id", "type": "bigint", "nullable": "NO", "key": "PRI"},
            {"name": "value", "type": "varchar(20)", "nullable": "YES", "key": ""},
        ]
        with patch.object(matrix_core, "fetch_columns", side_effect=[columns, columns]), \
                patch.object(matrix_core, "count_table", side_effect=[1, 1]), \
                patch.object(matrix_core, "fetch_table_sample",
                             side_effect=[(["id", "value"], [[1, "same"]]),
                                          (["id", "value"], [[1, "same"]])]) as fetch:
            result = matrix_core.compare_one_table(
                object(),
                {"cluster": "A", "db": "db", "table": "orders"},
                {"cluster": "B", "db": "db", "table": "orders"},
                keys=["t.id"],
            )
        self.assertEqual(["id"], fetch.call_args_list[0].args[5])
        self.assertEqual(["id"], result["data"]["stats"]["keys"])

    def test_file_compare_does_not_claim_structure_or_key_only_equivalence(self):
        structure = excel_diff.compare_tables(
            ["id", "old_value"], [["1", "same"]],
            ["id", "new_value"], [["1", "same"]], ["id"])
        self.assertEqual("diff", structure["verdict"]["level"])
        self.assertIn("字段结构不一致", structure["verdict"]["text"])

        key_only = excel_diff.compare_tables(
            ["id"], [["1"]], ["id"], [["1"]], ["id"])
        self.assertEqual("warn", key_only["verdict"]["level"])
        self.assertIn("未比较任何非主键", key_only["verdict"]["text"])

    def test_large_duplicate_groups_do_not_build_quadratic_cost_matrix(self):
        left = [[f"A{i}"] for i in range(250)]
        right = [[f"B{i}"] for i in range(250)]
        meta = {}
        with patch.object(excel_diff, "hungarian", side_effect=AssertionError), \
                patch.object(excel_diff, "greedy_assign", side_effect=AssertionError):
            pairs, only_a, only_b = excel_diff.match_group(left, right, meta)
        self.assertEqual(250, len(pairs))
        self.assertEqual([], only_a)
        self.assertEqual([], only_b)
        self.assertTrue(meta["approximate"])

    def test_time_upper_bounds_include_fractional_last_second(self):
        self.assertEqual("2026-08-17 23:59:59.999999",
                         sql_tools._as_hi("2026-08-17"))
        self.assertEqual("2026-08-17 12:34:59.999999",
                         sql_tools._as_hi("2026-08-17 12:34"))
        self.assertEqual("2026-08-17 12:34:56.999999",
                         sql_tools._as_hi("2026-08-17 12:34:56"))
        self.assertEqual("2026-08-17 23:59:59.999999",
                         sql_tools._seg_end("2026-08-17", "天"))

    def test_insert_rejects_duplicate_excel_headers(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "duplicate.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["id", "id", "amount"])
            ws.append(["A", "B", 10])
            wb.save(source)
            wb.close()
            with self.assertRaisesRegex(ValueError, "重复表头"):
                convert.excel_to_insert(
                    source,
                    "CREATE TABLE demo (`id` varchar(20), `amount` int)")

    def test_file_report_explicitly_marks_detail_truncation(self):
        result = excel_diff.compare_tables(
            ["id", "amount"], [["1", "A"], ["2", "A"]],
            ["id", "amount"], [["1", "B"], ["2", "B"]], ["id"])
        with patch.object(excel_diff, "MAX_EXCEL_DIFF_ROWS", 1):
            report = excel_diff.build_report(result)
        summary = dict(report["结论"][1])
        self.assertEqual("1/2", summary["字段差异明细（导出/总数）"])
        self.assertIn("截断", summary["导出完整性"])


if __name__ == "__main__":
    unittest.main()
